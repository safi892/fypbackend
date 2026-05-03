import re
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_file_path, docx_file_path):
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    lines = content.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        if line.startswith('# '):
            title = line[2:]
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        elif line.startswith('**') and '**' in line[2:]:
            parts = line.split('**')
            if len(parts) >= 3:
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        run = p.add_run(part)
                        run.bold = True
                    else:
                        p.add_run(part)
                continue
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            match = re.match(r'^(\d+\.\s)(.*)', line)
            if match:
                doc.add_paragraph(match.group(2), style='List Number')
        elif line.startswith('|'):
            continue
        elif line.startswith('---'):
            doc.add_paragraph('_' * 60)
        elif line.strip():
            doc.add_paragraph(line)

    doc.save(docx_file_path)
    print(f"✅ Converted: {md_file_path} → {docx_file_path}")

if __name__ == "__main__":
    md_files = [
        ("API_Explanation.md", "API_Explanation.docx"),
        ("API_Interview_Questions.md", "API_Interview_Questions.docx"),
        ("CodeT5_Interview_Questions.md", "CodeT5_Interview_Questions.docx"),
    ]

    for md_file, docx_file in md_files:
        md_path = os.path.join("/Volumes/Data/saffi/back", md_file)
        docx_path = os.path.join("/Volumes/Data/saffi/back", docx_file)
        
        if os.path.exists(md_path):
            markdown_to_docx(md_path, docx_path)
        else:
            print(f"⚠️ File not found: {md_path}")